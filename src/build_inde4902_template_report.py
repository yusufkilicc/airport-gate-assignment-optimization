from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"C:\Users\yusuf\Downloads\INDE4902_Project_Report_Template 2026.docx")
OUTPUT = ROOT / "reports" / "INDE4902_TIA_Project_Report.docx"
ASSET_DIR = ROOT / "reports" / "_inde4902_assets"
FIG_DIR = ROOT / "outputs" / "figures"

TITLE = (
    "Integrated Taxi and Apron Optimization Framework for Tirana International Airport (TIA): "
    "Gate Allocation and Turnaround Capacity Analysis"
)
STUDENT_NAME = "Yusuf Kılıç"
STUDENT_ID = "23INDE1049"
SUPERVISOR = "Sonya Javadi"


def clear_body(document: Document) -> None:
    body = document._body._element
    sect_pr = body.sectPr
    sect_clone = deepcopy(sect_pr) if sect_pr is not None else None
    for child in list(body):
        body.remove(child)
    if sect_clone is not None:
        body.append(sect_clone)


def configure_section(section) -> None:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.1)
    section.right_margin = Inches(1.1)


def prepare_template_logo() -> Path | None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    logo = ASSET_DIR / "isik_logo.png"
    try:
        with ZipFile(TEMPLATE) as zf:
            data = zf.read("word/media/image1.png")
        logo.write_bytes(data)
        return logo
    except Exception:
        return None


def get_style(document: Document, name: str, fallback: str = "Normal") -> str:
    try:
        document.styles[name]
        return name
    except KeyError:
        return fallback


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_para(
    document: Document,
    text: str = "",
    style: str = "Normal",
    align: int | None = None,
    space_after: float | None = None,
    bold: bool = False,
    italic: bool = False,
) -> None:
    p = document.add_paragraph(style=get_style(document, style))
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic


def add_body_paragraphs(document: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        add_para(document, text, "Normal", space_after=6)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        add_para(document, item, "List Bullet", space_after=3)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        add_para(document, item, "List Number", space_after=3)


def add_heading2(document: Document, text: str) -> None:
    add_para(document, text, "Heading 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def add_chapter(document: Document, title: str) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)
    add_para(document, title, "Heading 3", space_after=12)


def add_subheading(document: Document, title: str) -> None:
    add_para(document, title, "Heading 5", space_after=6)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def format_table(table, header_fill: str = "D9EAF7", font_size: float = 9.0) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
            if row_idx == 0:
                shade_cell(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_caption(document: Document, text: str, above: bool = False) -> None:
    p = document.add_paragraph(style=get_style(document, "Caption"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6 if above else 10)
    p.paragraph_format.keep_together = True
    if above:
        p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.italic = True


def add_compact_line(document: Document, text: str, indent: float = 0.0) -> None:
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.88
    run = p.add_run(text)
    run.font.size = Pt(8.2)


def add_table(document: Document, caption: str, headers: list[str], rows: list[list[str]], font_size: float = 8.8) -> None:
    add_caption(document, caption, above=True)
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cells[idx].text = value
    format_table(table, font_size=font_size)


def add_figure(document: Document, image_name: str, caption: str, width: float = 6.0) -> None:
    image_path = FIG_DIR / image_name
    if not image_path.exists():
        add_para(document, f"[Missing figure: {image_name}]", "Normal", WD_ALIGN_PARAGRAPH.CENTER, italic=True)
        add_caption(document, caption)
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(image_path), width=Inches(width))
    add_caption(document, caption)


def add_equation(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)


def add_case_note(document: Document) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    cell = table.cell(0, 0)
    shade_cell(cell, "FFF2CC")
    set_cell_margins(cell, top=140, start=140, bottom=140, end=140)
    cell.text = (
        "Case context note: Tirana International Airport (TIA, Albania) is used as the representative "
        "airport context for the report. The schedule, stand inventory, penalties, and scenario inputs "
        "are synthetic educational inputs designed to represent a TIA-like mid-sized international "
        "operation; they are not official or confidential TIA operating records."
    )
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.bold = True
    add_para(document, "", "Normal", space_after=6)


def add_front_matter(document: Document) -> None:
    logo = prepare_template_logo()
    if logo:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(1.25))

    add_para(document, "IŞIK UNIVERSITY", "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=2, bold=True)
    add_para(document, "Faculty of Natural Sciences and Engineering", "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(document, "Department of Industrial Engineering", "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(document, TITLE, "Title Page", WD_ALIGN_PARAGRAPH.CENTER, space_after=18, bold=True)
    add_para(document, "Project Report", "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(document, "by", "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(document, STUDENT_NAME, "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(document, STUDENT_ID, "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=22)
    add_para(document, f"Supervised by: {SUPERVISOR}", "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(document, "May 2026", "Title Page", WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    document.add_page_break()

    add_para(document, TITLE, "Title Page", WD_ALIGN_PARAGRAPH.CENTER, space_after=18, bold=True)
    add_para(document, "A Project Presented by", "Title Page", WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(document, STUDENT_NAME, "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(document, STUDENT_ID, "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    add_para(document, "to", "Title Page", WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(document, "Işık University Department of Industrial Engineering", "Title Page", WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(document, "in partial fulfillment of the requirements for the INDE4902 Project Report", "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(document, "May 2026", "Title Page", WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    document.add_page_break()

    add_para(document, "© Copyright 2026", "Title Page", WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(document, STUDENT_NAME, "Body Text 2", WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    document.add_page_break()

    add_heading2(document, "Abstract")
    add_body_paragraphs(
        document,
        [
            "This project develops an integrated taxi and apron decision-support framework for a representative Tirana International Airport (TIA, Albania) case, with the operational core focused on gate and stand assignment under peak-day congestion. Contact gates and remote stands are treated as scarce service resources, and poor assignments are evaluated through passenger walking burden, terminal-zone mismatch, remote-stand exposure, and turnaround friction.",
            "The analysis uses a synthetic but operationally consistent peak-day case with 36 flights, 7 contact gates, and 4 remote stands. Three policies are compared: first-come-first-served (FCFS), a passenger-priority heuristic called Priority Dispatch, and a binary mixed-integer linear programming (MILP) model solved with CBC through PuLP. The project also includes sensitivity analysis, a 100-scenario delay simulation, a demand stress test, figures, an interactive dashboard, and reproducible outputs.",
            "Compared with FCFS, the optimized MILP reduces total assignment cost by 14.70%, weighted average walking distance by 17.95%, zone-mismatch flights by 52.63%, narrow-body use of wide-body gates by 87.50%, and expected service extension by 6.79%. Remote-flight count stays fixed because peak concurrent demand reaches 11 aircraft while only 7 contact gates are available, but passenger exposure to remote operations falls from 3,014 to 2,149. The heavy-demand scenario becomes infeasible under the current stand inventory.",
            "Keywords: airport operations, apron optimization, gate assignment, mixed-integer linear programming, Tirana International Airport, turnaround capacity, robustness analysis, supply chain management",
        ],
    )
    document.add_page_break()

    add_heading2(document, "Acknowledgments")
    add_body_paragraphs(
        document,
        [
            f"I would like to express my sincere gratitude to {SUPERVISOR} for her guidance and for providing the academic framework that shaped this project. Her feedback helped connect the technical modeling work with broader industrial engineering themes such as capacity planning, service performance, and decision support.",
            "I would also like to thank Işık University Department of Industrial Engineering for providing the learning environment that supported the development of this report. The project benefited from the analytical foundation gained through coursework in optimization, operations management, and supply chain systems.",
        ],
    )
    document.add_page_break()

    add_heading2(document, "Contents")
    contents = [
        ("toc 3", "Chapter 1 Introduction"),
        ("toc 5", "1.1 Project Background"),
        ("toc 5", "1.2 Motivation and Problem Statement"),
        ("toc 5", "1.3 Project Objectives"),
        ("toc 5", "1.4 Relevance to United Nations SDGs"),
        ("toc 5", "1.5 Contributions"),
        ("toc 3", "Chapter 2 Problem Description"),
        ("toc 5", "2.1 Detailed Problem Definition"),
        ("toc 5", "2.2 Scope and Assumptions"),
        ("toc 5", "2.3 Limitations and Restrictions"),
        ("toc 3", "Chapter 3 Literature Review"),
        ("toc 5", "3.1 Review of Relevant Studies"),
        ("toc 5", "3.2 Summary and Taxonomy"),
        ("toc 3", "Chapter 4 System Design and Methodology"),
        ("toc 5", "4.1 Proposed Solution Approach"),
        ("toc 5", "4.2 System Design / Mathematical Modeling"),
        ("toc 5", "4.3 Data Collection and Preparation"),
        ("toc 5", "4.4 Solution Method"),
        ("toc 3", "Chapter 5 Computational Results and Analysis"),
        ("toc 5", "5.1 Experimental Design / Scenarios"),
        ("toc 5", "5.2 Analysis of Results"),
        ("toc 5", "5.3 Limitations"),
        ("toc 3", "Chapter 6 Conclusions and Discussions"),
        ("toc 5", "6.1 Cost-Benefit Analysis"),
        ("toc 5", "6.2 Economical, Social, Ethical and Environmental Impacts"),
        ("toc 5", "6.3 Conclusions and Future Work"),
        ("toc 3", "Bibliography"),
        ("toc 3", "Appendix"),
    ]
    for style, text in contents:
        add_compact_line(document, text, indent=0.0 if style == "toc 3" else 0.22)
    document.add_page_break()

    add_heading2(document, "List of Figures")
    figures = [
        "Figure 4-1 Peak-day stand demand profile",
        "Figure 4-2 Turnaround-time distribution by aircraft type and operating status",
        "Figure 5-1 Comparative KPI dashboard for FCFS, Priority Dispatch, and Optimized MILP",
        "Figure 5-2 Baseline FCFS gate schedule",
        "Figure 5-3 Priority Dispatch gate schedule",
        "Figure 5-4 Optimized MILP gate schedule",
        "Figure 5-5 Sensitivity analysis of MILP objective weights",
        "Figure 5-6 Robustness simulation under random delays",
        "Figure 5-7 Scenario comparison for feasible MILP cases",
    ]
    for fig in figures:
        add_para(document, fig, "table of figures", space_after=2)
    document.add_page_break()

    add_heading2(document, "List of Tables")
    tables = [
        "Table 2-1 Scope and modeling assumptions",
        "Table 3-1 Literature taxonomy for the gate assignment and turnaround planning problem",
        "Table 4-1 Main notation used in the MILP formulation",
        "Table 5-1 KPI comparison across the three assignment policies",
        "Table 5-2 Simulation summary across 100 stochastic delay scenarios",
        "Table 5-3 Scenario feasibility summary",
        "Table 5-4 Heavy-scenario capacity-relaxation experiment",
        "Table 6-1 Cost-benefit interpretation of the proposed framework",
    ]
    for tbl in tables:
        add_para(document, tbl, "table of figures", space_after=2)


def chapter_1(document: Document) -> None:
    add_chapter(document, "Introduction")
    add_subheading(document, "Project Background")
    add_body_paragraphs(
        document,
        [
            "Airport operations are service supply chains in which scarce physical capacity must be matched with time-sensitive demand while protecting passenger experience and operational reliability. Gates, remote stands, taxi-side movements, ground handling resources, baggage processes, and passenger flows are connected; a weak decision in one part of the system can create pressure in several downstream activities.",
            "This project studies that problem through the context of Tirana International Airport (TIA, Albania). TIA is used as the representative airport setting because the project is concerned with the type of planning challenge faced by a mid-sized international airport: banked demand, limited contact-gate capacity, mixed aircraft types, and the need to choose when remote stands must be used. The model is not presented as an official TIA operational audit; instead, it is a TIA-representative educational case built for transparent industrial engineering analysis.",
            "The technical focus is gate and stand assignment, supported by taxi-side and service penalties that approximate apron operational friction. This is a natural first layer of an integrated taxi-apron optimization framework because gate decisions determine where aircraft stop, how passengers move, and how much flexibility remains for later arrivals.",
        ],
    )
    add_case_note(document)

    add_subheading(document, "Motivation and Problem Statement")
    add_body_paragraphs(
        document,
        [
            "During peak traffic banks, contact gates become premium resources. If a high-passenger or long-turnaround flight is assigned to a remote stand while a lower-impact flight occupies a better contact gate, the airport may suffer longer walking distances, more bus operations, higher service extension, and lower robustness when delays occur. A first-come-first-served rule can be feasible, but it may consume scarce gates early and leave poor options for flights that arrive later.",
            "The central problem is therefore not simply whether every flight can receive a stand. The real planning question is how to allocate limited gate and apron resources so that unavoidable trade-offs are absorbed by the flights where they create the least system-wide damage. This requires a structured decision rule that can compare many feasible assignments at once rather than committing to local choices one flight at a time.",
        ],
    )

    add_subheading(document, "Project Objectives")
    add_body_paragraphs(
        document,
        [
            "The project aims to build and evaluate a reproducible decision-support framework for airport gate and apron-capacity planning under congested operating conditions. The objectives are measurable and tied directly to operational KPIs.",
        ],
    )
    add_bullets(
        document,
        [
            "Construct a realistic peak-day operating case with 36 flights, 7 contact gates, 4 remote stands, and repeated gate-capacity pressure.",
            "Compare a baseline FCFS policy, a passenger-priority heuristic, and an optimized MILP assignment policy.",
            "Evaluate assignment quality through interpretable KPIs such as total cost, remote flights, weighted walking distance, zone mismatches, wide-body gate misuse, and expected service extension.",
            "Test robustness through sensitivity analysis, stochastic delay simulation, and named operating scenarios.",
            "Present the results in a template-compliant project report with clear figure placement, captions, tables, and a reproducible output inventory.",
        ],
    )

    add_subheading(document, "Relevance to United Nations SDGs")
    add_body_paragraphs(
        document,
        [
            "The project is most closely related to SDG 9, Industry, Innovation and Infrastructure, because it develops an analytical tool for improving the use of airport infrastructure without assuming immediate physical expansion. It also relates to SDG 11, Sustainable Cities and Communities, because airport access and reliability affect regional mobility and service quality.",
            "The framework also has indirect relevance to SDG 12, Responsible Consumption and Production, and SDG 13, Climate Action. Better stand allocation can reduce unnecessary passenger bus movements, avoid avoidable service friction, and support more efficient use of existing assets. Because the current model uses calibrated operational penalties rather than measured emissions data, the environmental interpretation is qualitative rather than a direct carbon accounting result.",
        ],
    )

    add_subheading(document, "Contributions")
    add_body_paragraphs(
        document,
        [
            "The project contributes an integrated, reproducible, and managerially interpretable framework for airport gate assignment. Analytically, it demonstrates how a compact MILP formulation can outperform reactive assignment rules on multiple service indicators at the same time. Managerially, it translates optimization output into measures that airport planning teams can understand: passenger exposure to remote stands, walking burden, zone consistency, and feasibility under stress.",
            "A second contribution is the explicit connection between deterministic optimization and robustness testing. The report does not stop at one nominal solution; it evaluates how the assignment policies behave under parameter changes, random delays, and heavier demand. This makes the framework more useful as a planning tool rather than just a one-time optimization exercise.",
        ],
    )


def chapter_2(document: Document) -> None:
    add_chapter(document, "Problem Description")
    add_subheading(document, "Detailed Problem Definition")
    add_body_paragraphs(
        document,
        [
            "For a nontechnical stakeholder, the problem can be described as follows. On a busy airport day, many aircraft need places to park for a limited amount of time. Some aircraft can use only certain stands because of aircraft size or operational requirements. Some stands are better for passengers because they are closer to the terminal or better matched to a terminal zone. Remote stands remain useful, but they usually create extra walking, bus movement, and service handling burden.",
            "The planner must assign every flight to exactly one feasible stand. Two flights cannot use the same stand at overlapping times, and aircraft compatibility must be respected. Because there are only 7 contact gates and peak concurrent demand reaches 11 aircraft, some remote usage is unavoidable. The project therefore asks which flights should receive the scarce contact gates and which flights should be assigned to remote stands when the system is under pressure.",
            "The assignment problem is evaluated through a service-cost perspective. The cost of an assignment increases when passengers walk farther, when a flight uses a remote stand, when the assigned terminal zone does not match the preferred zone, or when scarce wide-body-compatible gates are consumed by narrow-body aircraft. The optimization model searches across all feasible assignments to minimize the total weighted cost.",
        ],
    )

    add_subheading(document, "Scope and Assumptions")
    add_body_paragraphs(
        document,
        [
            "The scope is tactical and planning-oriented. The report does not attempt to run a live airport control tower or real-time apron management system. Instead, it evaluates whether structured assignment logic improves the quality of a peak-day stand plan compared with simpler rules.",
        ],
    )
    add_table(
        document,
        "Table 2-1 Scope and modeling assumptions",
        ["Area", "Included in the project", "Main assumption"],
        [
            ["Airport context", "Representative TIA case", "Synthetic inputs are scaled to a mid-sized international operation."],
            ["Planning horizon", "One peak operating day", "The day is the tactical unit for assignment and evaluation."],
            ["Flights", "36 scheduled flights", "Arrival and departure windows are known before assignment."],
            ["Stands", "7 contact gates and 4 remote stands", "Remote stands are feasible but less desirable from a service perspective."],
            ["Compatibility", "Aircraft size and stand capability", "Wide-body flights require compatible stands."],
            ["Uncertainty", "Delay simulation and named disruption scenario", "The core MILP is deterministic; uncertainty is tested externally."],
            ["Cost data", "Calibrated service penalties", "Penalties represent operational trade-offs, not audited financial costs."],
        ],
    )

    add_subheading(document, "Limitations and Restrictions")
    add_body_paragraphs(
        document,
        [
            "The main restriction is data availability. Airport stand assignments, airline preferences, and detailed taxi/apron movement records are usually operationally sensitive. For that reason, this project uses a synthetic case rather than confidential TIA records. This is appropriate for the academic purpose of demonstrating modeling logic, but it limits direct operational deployment.",
            "The current model focuses on stand assignment and includes taxi-side inconvenience as a cost component, but it does not explicitly optimize taxiway routing, pushback sequencing, runway queues, or ground-handling crew schedules. These elements are important in a complete airport operations model and are identified as future extensions.",
            "Finally, the cost coefficients are calibrated to represent relative service preferences. They are useful for comparing policies inside the same modeling environment, but they should not be interpreted as direct euro-denominated savings without further data collection.",
        ],
    )


def chapter_3(document: Document) -> None:
    add_chapter(document, "Literature Review")
    add_subheading(document, "Review of Relevant Studies")
    add_body_paragraphs(
        document,
        [
            "The supply chain perspective emphasized by Chopra and Meindl (2016) highlights that effective planning requires balancing capacity, responsiveness, and service performance rather than optimizing isolated tasks. This logic transfers naturally to airport environments, where one gate assignment can affect passengers, baggage, ground handling, apron utilization, and downstream operational flexibility.",
            "Gate assignment has been studied extensively in operations research. Mangoubi and Mathaisel (1985) presented one of the early optimization formulations for assigning flights to gates. Later work expanded the problem by incorporating time-overlap constraints, passenger walking distance, and congestion effects under dense schedules. Dorndorf, Drexl, Nikulin, and Pesch (2005) reviewed the field and showed that realistic gate scheduling problems combine feasibility, compatibility, and temporal-conflict constraints.",
            "Ding, Lim, Rodrigues, and Zhu (2004) are especially relevant because they examine over-constrained flight-to-gate assignments. Their work supports the managerial logic of this project: when premium gate capacity is insufficient for all overlapping flights, the key question becomes how to prioritize access to scarce high-quality resources.",
            "Uncertainty is also central in airport operations. Yan and Tang (2007) study airport gate assignment under stochastic delays, showing that robust planning is necessary when flight schedules shift. Das, Gzara, and Stützle (2020) review single-objective and multi-objective gate assignment approaches and confirm that passenger-oriented performance measures remain important in recent research. The broader turnaround literature, including Wu and Caves (2000) and Schmidt (2017), further links turnaround coordination to operational cost and reliability.",
        ],
    )

    add_subheading(document, "Summary and Taxonomy")
    add_body_paragraphs(
        document,
        [
            "Table 3-1 summarizes the literature streams most relevant to this project. The taxonomy shows that the proposed framework combines ideas from supply chain capacity planning, gate assignment optimization, stochastic robustness, and turnaround operations.",
        ],
    )
    add_table(
        document,
        "Table 3-1 Literature taxonomy for the gate assignment and turnaround planning problem",
        ["Source", "Main focus", "Method / perspective", "Relevance to this project"],
        [
            ["Chopra & Meindl (2016)", "Capacity, responsiveness, and supply chain trade-offs", "Supply chain management framework", "Supports the service supply chain interpretation of airport gate capacity."],
            ["Mangoubi & Mathaisel (1985)", "Flight-to-gate optimization", "Mathematical programming", "Provides early foundation for gate assignment as an optimization problem."],
            ["Ding et al. (2004)", "Over-constrained gate assignment", "Heuristics and optimization", "Matches the scarce-contact-gate logic of the TIA-representative case."],
            ["Dorndorf et al. (2005)", "State of the art in gate scheduling", "Literature survey", "Clarifies compatibility and temporal conflict requirements."],
            ["Yan & Tang (2007)", "Gate assignment under delays", "Stochastic / heuristic analysis", "Motivates the 100-scenario robustness simulation."],
            ["Das et al. (2020)", "Single vs. multi-objective gate assignment", "Review article", "Supports use of passenger-centric and multi-objective KPIs."],
            ["Wu & Caves (2000); Schmidt (2017)", "Aircraft turnaround performance", "Operational cost and simulation review", "Links stand assignment quality to turnaround efficiency and uncertainty."],
        ],
        font_size=7.4,
    )


def chapter_4(document: Document) -> None:
    add_chapter(document, "System Design and Methodology")
    add_subheading(document, "Proposed Solution Approach")
    add_body_paragraphs(
        document,
        [
            "The proposed framework has four analytical layers. First, a peak-day airport case is constructed with flight windows, passenger loads, aircraft categories, contact-gate characteristics, and remote-stand alternatives. Second, three assignment policies are applied to the same operating day: FCFS, Priority Dispatch, and Optimized MILP. Third, assignment outputs are converted into operational KPIs. Fourth, the results are tested through sensitivity analysis, delay simulation, and scenario analysis.",
            "This structure was selected because a single optimization run is not enough for a practical planning report. A useful decision-support framework must explain the baseline, show how improvement is achieved, and test whether the conclusion survives changes in assumptions. The method therefore combines mathematical optimization with transparent benchmarking and robustness analysis.",
        ],
    )

    add_subheading(document, "System Design / Mathematical Modeling")
    add_body_paragraphs(
        document,
        [
            "The core model is a binary mixed-integer linear program. It assigns each flight to one stand while respecting aircraft compatibility and preventing overlapping flights from sharing a stand. The objective minimizes a weighted assignment cost that combines walking burden, taxi-side inconvenience, remote-stand use, terminal-zone mismatch, and scarcity penalties for using wide-body-capable gates with narrow-body aircraft.",
        ],
    )
    add_table(
        document,
        "Table 4-1 Main notation used in the MILP formulation",
        ["Symbol", "Meaning", "Unit / type"],
        [
            ["F", "Set of flights", "Index set"],
            ["G", "Set of contact gates and remote stands", "Index set"],
            ["C", "Set of overlapping flight pairs", "Index set"],
            ["p_f", "Passenger load of flight f", "passengers"],
            ["w_g", "Walking distance associated with stand g", "meters"],
            ["c_fg", "Assignment cost of assigning flight f to stand g", "cost units"],
            ["x_fg", "1 if flight f is assigned to stand g; 0 otherwise", "binary variable"],
        ],
        font_size=7.8,
    )
    add_equation(document, "x_fg = 1 if flight f is assigned to stand g; 0 otherwise.                         (4.1)")
    add_equation(document, "min Z = sum_{f in F} sum_{g in G} c_fg x_fg                                      (4.2)")
    add_body_paragraphs(
        document,
        [
            "The objective function in (4.2) minimizes total assignment cost. The coefficient c_fg is calculated before solving the model and represents the combined operational burden of a flight-stand match.",
        ],
    )
    add_equation(document, "sum_{g in G} x_fg = 1                  for all f in F                             (4.3)")
    add_equation(document, "x_fg = 0                              for incompatible pairs (f,g)               (4.4)")
    add_equation(document, "x_ig + x_kg <= 1                      for all (i,k) in C and g in G              (4.5)")
    add_body_paragraphs(
        document,
        [
            "Constraint (4.3) assigns every flight to exactly one stand. Constraint (4.4) removes infeasible aircraft-stand combinations. Constraint (4.5) prevents temporal conflicts by ensuring that two overlapping flights cannot occupy the same stand at the same time.",
        ],
    )

    add_subheading(document, "Data Collection and Preparation")
    add_body_paragraphs(
        document,
        [
            "The case data is generated from the project pipeline rather than collected from confidential airport records. The schedule is designed to represent a busy TIA-like operating day with three traffic banks: 14 flights in the morning bank, 10 in the midday bank, and 12 in the evening bank. The fleet includes 26 narrow-body and 10 wide-body flights, and 47.22% of flights require international-processing capability.",
            "The stand system contains 11 total positions: 7 contact gates and 4 remote stands. Contact gates differ by walking distance, terminal zone, and aircraft compatibility. Remote stands are modeled as operationally feasible for the aircraft types in the case, but they impose longer walking distance and additional service burden.",
            "Figure 4-1 should be placed immediately after this paragraph because it establishes the structural capacity imbalance of the case before the optimization model is evaluated.",
        ],
    )
    add_figure(
        document,
        "capacity_profile.png",
        "Figure 4-1 Peak-day stand demand profile. The shaded regions indicate periods in which concurrent stand occupancy exceeds available contact-gate capacity, making remote-stand use operationally unavoidable.",
        width=6.2,
    )
    add_body_paragraphs(
        document,
        [
            "The second descriptive figure belongs directly after the discussion of flight service times. It shows why aircraft type and operating status matter: wide-body and international flights tend to create longer stand-occupancy windows, which increases overlap pressure during traffic banks.",
        ],
    )
    add_figure(
        document,
        "turnaround_distribution.png",
        "Figure 4-2 Turnaround-time distribution by aircraft type and operating status. Longer and more variable service windows intensify scheduling pressure during peak periods.",
        width=5.8,
    )

    add_subheading(document, "Solution Method")
    add_body_paragraphs(
        document,
        [
            "The implementation uses pandas and numpy for data handling, PuLP with the CBC solver for MILP optimization, matplotlib and seaborn for static figures, and Plotly Dash for the interactive dashboard. The full workflow is orchestrated through src/run_project.py, which rebuilds the processed data, assignment tables, KPI summaries, simulation outputs, scenario outputs, and figures.",
            "The three compared assignment policies are: FCFS, which processes flights chronologically; Priority Dispatch, which assigns higher-passenger flights before lower-passenger flights; and Optimized MILP, which chooses all assignments simultaneously to minimize total cost. Metrics are calculated after assignments are produced so the outputs remain interpretable in operational terms.",
        ],
    )


def chapter_5(document: Document) -> None:
    add_chapter(document, "Computational Results and Analysis")
    add_subheading(document, "Experimental Design / Scenarios")
    add_body_paragraphs(
        document,
        [
            "The computational analysis is organized into four experiment groups. The first compares FCFS, Priority Dispatch, and Optimized MILP on the base 36-flight day. The second visualizes the resulting gate schedules. The third performs sensitivity analysis on the MILP objective weights. The fourth evaluates robustness through a 100-scenario stochastic delay simulation and three named operating scenarios: normal, heavy, and disruption.",
            "The heavy scenario adds two flights to each traffic bank, increasing the day from 36 to 42 flights. This scenario is used as a stress test because it asks whether the current stand inventory can still produce a feasible assignment when demand grows. The disruption scenario applies arrival delays to approximately 35% of flights before assignment.",
        ],
    )

    add_subheading(document, "Analysis of Results")
    add_body_paragraphs(
        document,
        [
            "The KPI comparison in Table 5-1 shows that both structured policies improve on FCFS, while the MILP provides the strongest overall result.",
        ],
    )
    add_table(
        document,
        "Table 5-1 KPI comparison across the three assignment policies",
        ["Metric", "Baseline FCFS", "Priority Dispatch", "Optimized MILP"],
        [
            ["Total assignment cost", "66,649.0", "59,514.8", "56,850.0"],
            ["Remote flights", "11", "12", "11"],
            ["Contact-gate share (%)", "69.44", "66.67", "69.44"],
            ["Weighted average walking distance (m)", "550.19", "459.87", "451.44"],
            ["Weighted passenger distance (km)", "4,010.90", "3,352.48", "3,291.00"],
            ["Zone mismatch flights", "19", "14", "9"],
            ["Narrow-body on wide-body gates", "8", "0", "1"],
            ["Total expected service extension (min)", "560", "569", "522"],
            ["Maximum gate utilization (%)", "53.41", "50.27", "51.68"],
        ],
        font_size=8.2,
    )
    add_body_paragraphs(
        document,
        [
            "Relative to FCFS, Priority Dispatch lowers total assignment cost by 10.70% and reduces weighted walking by 16.42%. It also eliminates narrow-body use of wide-body gates, which is important for protecting scarce high-capability resources. However, its greedy passenger-first structure sends 12 flights to remote stands and produces a slightly higher service-extension total than FCFS.",
            "Relative to FCFS, the Optimized MILP reduces total assignment cost by 14.70%, lowers weighted average walking distance by 17.95%, decreases zone-mismatch flights by 52.63%, reduces narrow-body use of wide-body gates by 87.50%, and lowers total expected service extension by 6.79%. The number of remote flights remains 11 because remote usage is structurally required by the peak-demand pattern, but the passengers exposed to remote operations fall from 3,014 to 2,149.",
        ],
    )
    add_figure(
        document,
        "kpi_comparison.png",
        "Figure 5-1 Comparative KPI dashboard for FCFS, Priority Dispatch, and Optimized MILP. The figure should be placed immediately after the KPI table and interpretation so the reader can connect the numerical and visual evidence.",
        width=6.0,
    )
    add_body_paragraphs(
        document,
        [
            "The gate-schedule figures should follow the KPI comparison because they explain how the aggregate performance differences arise. FCFS is reactive and may consume premium gates early. Priority Dispatch is more intentional but still makes local choices. The MILP schedule is the most disciplined because it coordinates assignments across the full operating day.",
        ],
    )
    add_figure(
        document,
        "baseline_gantt.png",
        "Figure 5-2 Baseline FCFS gate schedule. Premium stands are consumed early and flexibility is reduced for later flights.",
        width=6.2,
    )
    add_figure(
        document,
        "priority_gantt.png",
        "Figure 5-3 Priority Dispatch gate schedule. High-passenger flights receive better positions, but the heuristic still makes myopic local decisions.",
        width=6.2,
    )
    add_figure(
        document,
        "optimized_gantt.png",
        "Figure 5-4 Optimized MILP gate schedule. The solution protects scarce resources more deliberately across the operating day.",
        width=6.2,
    )
    add_body_paragraphs(
        document,
        [
            "The sensitivity analysis tests whether the main conclusion depends on one particular cost calibration. The remote-stand penalty changes total cost but does not change the number of remote flights in the tested range, confirming that remote usage is primarily driven by physical capacity. Zone and scarcity penalties have stronger behavioral effects because they change how contact gates are distributed among competing flights.",
        ],
    )
    add_figure(
        document,
        "sensitivity_analysis.png",
        "Figure 5-5 Sensitivity analysis of MILP objective weights. Remote usage remains structurally constrained, while zone-mismatch and gate-scarcity penalties materially affect assignment behavior.",
        width=6.0,
    )
    add_body_paragraphs(
        document,
        [
            "The delay simulation results in Table 5-2 show that the optimized policy remains strongest under uncertainty. Mean KPI values are computed over feasible scenarios only, while feasibility rate is reported separately as a reliability measure.",
        ],
    )
    add_table(
        document,
        "Table 5-2 Simulation summary across 100 stochastic delay scenarios",
        ["Method", "Mean total cost", "Std. dev.", "Mean remote flights", "Mean zone mismatches", "Mean weighted walk (m)", "Feasibility rate"],
        [
            ["FCFS", "61,587.27", "3,407.26", "9.74", "18.03", "509.30", "0.80"],
            ["Priority Dispatch", "53,929.33", "2,894.69", "9.83", "15.18", "424.15", "0.78"],
            ["Optimized MILP", "50,168.71", "3,001.23", "8.52", "11.51", "405.16", "0.81"],
        ],
        font_size=7.7,
    )
    add_body_paragraphs(
        document,
        [
            "Compared with FCFS, MILP improves mean total assignment cost by 18.54%, mean zone mismatches by 36.17%, and mean weighted walking distance by 20.45%. Priority Dispatch also remains better than FCFS, but its feasibility rate is slightly lower than both FCFS and MILP. The key interpretation is that optimization does not only improve the nominal plan; it also degrades more gracefully when delays occur.",
        ],
    )
    add_figure(
        document,
        "simulation_results.png",
        "Figure 5-6 Robustness simulation under random delays. MILP remains strongest on the main service-oriented metrics, while Priority Dispatch stays between MILP and FCFS.",
        width=6.0,
    )
    add_body_paragraphs(
        document,
        [
            "The named scenario analysis separates ordinary variation from structural capacity failure. The normal and disruption scenarios are feasible, but the heavy scenario is infeasible under both FCFS and MILP. This means the 42-flight demand level exceeds the modeled stand system, so the limitation is infrastructure capacity rather than assignment-policy quality.",
        ],
    )
    add_table(
        document,
        "Table 5-3 Scenario feasibility summary",
        ["Scenario", "Baseline FCFS status", "Optimized MILP status", "Interpretation"],
        [
            ["Normal", "Feasible", "Feasible", "Standard comparative case"],
            ["Heavy", "Infeasible", "Infeasible", "Current stand inventory is insufficient"],
            ["Disruption", "Feasible", "Feasible", "Optimization still improves service outcomes"],
        ],
        font_size=8.4,
    )
    add_table(
        document,
        "Table 5-4 Heavy-scenario capacity-relaxation experiment",
        ["Added remote stands", "Feasibility", "Remote flights", "Cost shift vs. normal"],
        [
            ["+0 (current)", "Infeasible", "-", "-"],
            ["+1", "Infeasible", "-", "-"],
            ["+2", "Infeasible", "-", "-"],
            ["+3", "Feasible", "15", "+23.2% vs. normal"],
        ],
        font_size=8.4,
    )
    add_body_paragraphs(
        document,
        [
            "The relaxation experiment shows that the heavy-demand infeasibility is structural rather than a modeling artifact. Adding one or two remote stands is not enough; three additional remote stands are required before the heavy schedule becomes feasible. Even then, the optimal plan places 15 flights on remote stands, indicating a meaningful capacity gap.",
        ],
    )
    add_figure(
        document,
        "scenario_comparison.png",
        "Figure 5-7 Scenario comparison for feasible MILP cases. The heavy scenario is omitted because the modeled airport configuration becomes infeasible at that demand level.",
        width=5.9,
    )

    add_subheading(document, "Limitations")
    add_body_paragraphs(
        document,
        [
            "The computational results should be interpreted within the assumptions of the synthetic case. The model captures gate and stand assignment quality, but it does not include full taxiway routing, runway queues, gate towing, crew constraints, or real-time recovery actions. The financial interpretation is also limited because the cost coefficients are calibrated planning penalties rather than observed accounting data.",
            "These limitations do not undermine the main conclusion because the report is designed as an academic decision-support framework. They do, however, define the boundary between this project and a production airport operations system.",
        ],
    )


def chapter_6(document: Document) -> None:
    add_chapter(document, "Conclusions and Discussions")
    add_subheading(document, "Cost-Benefit Analysis")
    add_body_paragraphs(
        document,
        [
            "Because the project uses synthetic operational inputs, the cost-benefit analysis is expressed in planning terms rather than audited financial savings. The benefits are measured through reductions in weighted walking burden, zone mismatches, wide-body gate misuse, service extension, and assignment cost. These indicators represent operational value that could later be converted into financial estimates if actual TIA data became available.",
        ],
    )
    add_table(
        document,
        "Table 6-1 Cost-benefit interpretation of the proposed framework",
        ["Benefit area", "Observed project result", "Managerial interpretation"],
        [
            ["Passenger experience", "17.95% lower weighted average walking distance vs. FCFS", "Fewer passengers are exposed to poor stand choices."],
            ["Remote-stand exposure", "Remote passenger exposure falls from 3,014 to 2,149", "The same structural remote count can be allocated more intelligently."],
            ["Gate compatibility", "Narrow-body use of wide-body gates falls from 8 to 1", "Scarce high-capability gates are protected for aircraft that need them."],
            ["Operational robustness", "MILP remains best under 100 delay scenarios", "The plan is less fragile under schedule uncertainty."],
            ["Implementation cost", "Uses reproducible Python pipeline and open-source CBC solver", "A prototype can be maintained without expensive solver dependency at this scale."],
        ],
        font_size=8.2,
    )
    add_body_paragraphs(
        document,
        [
            "The main implementation cost would be data integration and operational validation. In a real airport setting, the model would need reliable schedule feeds, stand availability data, compatibility rules, airline preferences, and agreement on penalty calibration. The prototype shows that the analytical gain is large enough to justify further validation if the airport wants a lightweight planning tool.",
        ],
    )

    add_subheading(document, "Economical, Social, Ethical and Environmental Impacts")
    add_body_paragraphs(
        document,
        [
            "Economically, the framework supports better use of existing infrastructure. Instead of immediately assuming new gates are required, the model identifies how much improvement can be achieved through smarter allocation and when demand becomes structurally infeasible. This distinction helps planners separate operational improvement from long-term capital investment needs.",
            "Socially, the model improves passenger-oriented service indicators such as walking burden and terminal-zone consistency. Ethical considerations mainly concern transparency and data use. If real airport data is later used, the project should protect sensitive operational records and avoid presenting synthetic results as official airport performance.",
            "Environmentally, better assignment can reduce avoidable remote-stand exposure and related service movements, but this report does not calculate direct emissions. Future work should integrate taxi fuel burn, bus operations, and ground-service emissions if environmental impact is to be quantified formally.",
        ],
    )

    add_subheading(document, "Conclusions and Future Work")
    add_body_paragraphs(
        document,
        [
            "This project shows that airport turnaround and apron-capacity planning can be analyzed effectively through a supply chain perspective. Gates and remote stands operate as scarce service resources, and the quality of their allocation affects passenger burden, operational friction, and future flexibility.",
            "The central finding is that structured assignment logic creates measurable value even under tight capacity. Priority Dispatch closes much of the gap between naive dispatching and optimization, while the MILP delivers the strongest overall performance in the base case, under random delays, and in feasible named scenarios. The heavy-demand stress test also shows that infrastructure limits dominate policy quality once demand exceeds feasible stand capacity.",
            "Future work should extend the model with real airport data if access becomes available, integrate taxiway routing and pushback conflicts, include ground-handling resource constraints, estimate direct financial and environmental impacts, and develop a dynamic recovery model for disruptions after the operating day has begun.",
        ],
    )


def bibliography_and_appendix(document: Document) -> None:
    add_chapter(document, "Bibliography")
    references = [
        "Chopra, S., & Meindl, P. (2016). Supply chain management: Strategy, planning, and operation (6th ed.). Pearson.",
        "Das, G. S., Gzara, F., & Stützle, T. (2020). A review on airport gate assignment problems: Single versus multi-objective approaches. Omega, 92, 102146. https://doi.org/10.1016/j.omega.2019.102146",
        "Ding, H., Lim, A., Rodrigues, B., & Zhu, Y. (2004). New heuristics for over-constrained flight-to-gate assignments. Journal of the Operational Research Society, 55(7), 760-768. https://doi.org/10.1057/palgrave.jors.2601736",
        "Dorndorf, U., Drexl, A., Nikulin, Y., & Pesch, E. (2005). Flight gate scheduling: State-of-the-art and recent developments (Working Paper No. 584). Christian-Albrechts-Universität zu Kiel.",
        "Mangoubi, R. S., & Mathaisel, D. F. X. (1985). Optimizing gate assignments at airport terminals. Transportation Science, 19(2), 173-188. https://doi.org/10.1287/trsc.19.2.173",
        "Schmidt, M. (2017). A review of aircraft turnaround operations and simulations. Progress in Aerospace Sciences, 92, 25-38. https://doi.org/10.1016/j.paerosci.2017.05.002",
        "Wu, C.-L., & Caves, R. E. (2000). Aircraft operational costs and turnaround efficiency at airports. Journal of Air Transport Management, 6(4), 201-208. https://doi.org/10.1016/S0969-6997(00)00014-4",
        "Yan, S., & Tang, C.-H. (2007). A heuristic approach for airport gate assignments under stochastic flight delays. European Journal of Operational Research, 180(2), 547-567.",
    ]
    for ref in references:
        add_para(document, ref, "Normal", space_after=6)

    add_chapter(document, "Appendix")
    add_subheading(document, "Appendix A: Reproducible Output Inventory")
    add_body_paragraphs(
        document,
        [
            "The main reproducible outputs generated by the project are listed below. They should be submitted together with the report if the instructor requests supporting artifacts.",
        ],
    )
    add_bullets(
        document,
        [
            "data/processed/flights.csv",
            "data/processed/gates.csv",
            "data/processed/baseline_assignments.csv",
            "data/processed/priority_assignments.csv",
            "data/processed/optimized_assignments.csv",
            "outputs/kpi_summary.csv",
            "outputs/sensitivity_results.csv",
            "outputs/simulation_results.csv",
            "outputs/simulation_summary.csv",
            "outputs/scenario_kpi.csv",
            "outputs/project_summary.md",
            "outputs/figures/capacity_profile.png",
            "outputs/figures/turnaround_distribution.png",
            "outputs/figures/kpi_comparison.png",
            "outputs/figures/baseline_gantt.png",
            "outputs/figures/priority_gantt.png",
            "outputs/figures/optimized_gantt.png",
            "outputs/figures/sensitivity_analysis.png",
            "outputs/figures/simulation_results.png",
            "outputs/figures/scenario_comparison.png",
        ],
    )
    add_subheading(document, "Appendix B: Visual Placement Guide")
    add_body_paragraphs(
        document,
        [
            "The report already places the required visuals in the recommended locations. The table below documents exactly where each figure belongs and why it is placed there.",
        ],
    )
    document.add_page_break()
    add_table(
        document,
        "Table A-1 Detailed figure placement guide",
        ["Figure", "Source file", "Recommended placement", "Purpose"],
        [
            ["Figure 4-1", "capacity_profile.png", "Chapter 4, Data Collection and Preparation, after the paragraph introducing peak concurrent demand", "Shows why remote stands are structurally unavoidable."],
            ["Figure 4-2", "turnaround_distribution.png", "Chapter 4, after the paragraph discussing service-time variation", "Explains why aircraft type and operating status affect overlap pressure."],
            ["Figure 5-1", "kpi_comparison.png", "Chapter 5, immediately after Table 5-1 and the KPI interpretation", "Summarizes comparative performance visually."],
            ["Figure 5-2", "baseline_gantt.png", "Chapter 5, start of gate-schedule comparison", "Shows the reactive FCFS assignment pattern."],
            ["Figure 5-3", "priority_gantt.png", "Chapter 5, after Figure 5-2", "Shows how passenger-priority sequencing changes assignments."],
            ["Figure 5-4", "optimized_gantt.png", "Chapter 5, after Figure 5-3", "Shows the globally coordinated MILP schedule."],
            ["Figure 5-5", "sensitivity_analysis.png", "Chapter 5, after sensitivity-analysis explanation", "Tests whether conclusions depend on one calibration."],
            ["Figure 5-6", "simulation_results.png", "Chapter 5, after simulation table and interpretation", "Shows performance distribution under 100 random delay scenarios."],
            ["Figure 5-7", "scenario_comparison.png", "Chapter 5, after scenario feasibility tables", "Compares feasible scenario outcomes and highlights heavy-case infeasibility."],
        ],
        font_size=6.4,
    )


def apply_document_defaults(document: Document) -> None:
    for section in document.sections:
        configure_section(section)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    for style_name in ["Heading 2", "Heading 3", "Heading 5", "Caption"]:
        try:
            style = document.styles[style_name]
            style.font.name = "Times New Roman"
        except KeyError:
            continue
    try:
        document.styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)
        document.styles["Heading 3"].font.color.rgb = RGBColor(31, 78, 121)
        document.styles["Heading 5"].font.color.rgb = RGBColor(31, 78, 121)
    except KeyError:
        pass


def trim_trailing_empty_paragraphs(document: Document) -> None:
    body = document._body._element
    children = list(body)
    for child in reversed(children):
        if child.tag == qn("w:sectPr"):
            continue
        if child.tag != qn("w:p"):
            break
        texts = child.findall(".//" + qn("w:t"))
        has_text = any((node.text or "").strip() for node in texts)
        drawings = child.findall(".//" + qn("w:drawing"))
        if has_text or drawings:
            break
        body.remove(child)


def build_report() -> Path:
    if not TEMPLATE.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE}")
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    document = Document(TEMPLATE)
    clear_body(document)
    apply_document_defaults(document)

    add_front_matter(document)
    chapter_1(document)
    chapter_2(document)
    chapter_3(document)
    chapter_4(document)
    chapter_5(document)
    chapter_6(document)
    bibliography_and_appendix(document)

    document.core_properties.title = TITLE
    document.core_properties.author = STUDENT_NAME
    document.core_properties.subject = "INDE4902 Project Report"
    document.core_properties.keywords = "airport operations; gate assignment; MILP; TIA; apron optimization"
    trim_trailing_empty_paragraphs(document)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_report()
    print(f"Created {path}")
