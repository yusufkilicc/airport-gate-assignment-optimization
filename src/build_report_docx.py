from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


HEADING_RE = re.compile(r"^(#{1,3})\s+(.*)$")
NUMBERED_RE = re.compile(r"^\d+\.\s+(.*)$")
IMAGE_RE = re.compile(r"^!\[(.*?)\]\((.*?)\)$")


def add_table(document: Document, table_lines: list[str]) -> None:
    rows = []
    for line in table_lines:
        if not line.strip():
            continue
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        rows.append(cells)
    if len(rows) < 2:
        return

    header = rows[0]
    data_rows = rows[2:]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    for idx, value in enumerate(header):
        table.rows[0].cells[idx].text = value

    for row_values in data_rows:
        row = table.add_row().cells
        for idx, value in enumerate(row_values):
            row[idx].text = value


def build_docx() -> Path:
    root = Path(__file__).resolve().parents[1]
    report_md = root / "reports" / "INDE4313_final_report.md"
    output_docx = root / "reports" / "INDE4313_final_report.docx"

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = report_md.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue

        if in_code_block:
            p = document.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table(document, table_lines)
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            image_path = (report_md.parent / image_match.group(2)).resolve()
            if image_path.exists():
                document.add_picture(str(image_path), width=Inches(6.3))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            document.add_heading(text, level=level)
            i += 1
            continue

        numbered_match = NUMBERED_RE.match(stripped)
        if numbered_match:
            document.add_paragraph(numbered_match.group(1), style="List Number")
            i += 1
            continue

        if stripped.startswith("- "):
            document.add_paragraph(stripped[2:], style="List Bullet")
            i += 1
            continue

        if stripped.startswith("> "):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            run = p.add_run(stripped[2:])
            run.italic = True
            i += 1
            continue

        if stripped.startswith("*Figure") and stripped.endswith("*"):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            i += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.add_run(stripped)
        i += 1

    document.save(output_docx)
    return output_docx


if __name__ == "__main__":
    path = build_docx()
    print(f"Report created: {path}")
